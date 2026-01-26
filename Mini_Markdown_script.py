#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
import sys
import os
import shutil
from dataclasses import dataclass
from pathlib import Path

from PySide6.QtCore import Qt, QTimer
from PySide6.QtGui import QTextCursor
from PySide6.QtGui import QAction, QKeySequence, QFont, QTextDocument
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QFileDialog, QMessageBox,
    QPlainTextEdit, QTextEdit, QSplitter
)
from PySide6.QtPrintSupport import QPrinter


# -----------------------------
# DOCX export (Markdown -> docx)
# -----------------------------

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", re.DOTALL)

def _add_inlines_docx(paragraph, text: str):
    """
    Minimal inline markdown -> docx runs:
    **bold**, *italic*, `code`
    (Pas de gestion d’imbrication complexe : volontairement simple.)
    """
    from docx.shared import Pt  # lazy import

    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)

def export_docx_from_markdown(md: str, out_path: Path):
    """
    Convertisseur volontairement “sobre” :
    - Titres #..######
    - Listes -/* et 1.
    - Paragraphes
    - Blocs de code ``` ... ```
    - Inline ** * ``
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    in_code = False
    code_lines = []

    lines = md.replace("\r\n", "\n").replace("\r", "\n").split("\n")

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        # Un bloc de code simple (une série de paragraphes monospace)
        for cl in code_lines:
            p = doc.add_paragraph()
            run = p.add_run(cl)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        code_lines = []

    for line in lines:
        # Code fence
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                flush_code()
            continue

        if in_code:
            code_lines.append(line)
            continue

        # Titres
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            doc.add_heading(text, level=level)
            continue

        # Listes
        m_bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m_bullet:
            p = doc.add_paragraph(style="List Bullet")
            _add_inlines_docx(p, m_bullet.group(1).strip())
            continue

        m_num = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            _add_inlines_docx(p, m_num.group(1).strip())
            continue

        # Ligne vide => séparation (Word gère assez bien sans forcer)
        if not line.strip():
            doc.add_paragraph("")
            continue

        # Paragraphe normal
        p = doc.add_paragraph()
        _add_inlines_docx(p, line)

    # Si code non fermé
    if in_code:
        flush_code()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# -----------------------------
# Preview (éditable) + autosave
# -----------------------------

class PreviewEdit(QTextEdit):
    """
    Aperçu rendu MAIS éditable (pour couper/copier/coller facilement).
    La mise à jour depuis le Markdown est suspendue quand le focus est à droite,
    pour ne pas écraser tes collages.
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self.on_focus_in = None
        self.on_focus_out = None

        self.setTextInteractionFlags(
            Qt.TextEditorInteraction | Qt.TextSelectableByMouse | Qt.LinksAccessibleByMouse
        )

    def focusInEvent(self, event):
        if callable(self.on_focus_in):
            self.on_focus_in()
        super().focusInEvent(event)

    def focusOutEvent(self, event):
        super().focusOutEvent(event)
        if callable(self.on_focus_out):
            self.on_focus_out()


@dataclass
class AutosaveConfig:
    enabled: bool = True
    idle_ms: int = 1000            # autosave 1s après la dernière frappe
    use_main_file_if_possible: bool = True  # si un fichier est ouvert/sauvé, autosave dessus
    fallback_filename: str = "MiniMarkdown_autosave.md"  # si aucun fichier courant


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Mini Markdown — split editor")

        self.current_path: Path | None = None
        self.cfg = AutosaveConfig()

        self.pandoc_path = self._find_pandoc()
        self.has_pandoc = self.pandoc_path is not None

        # Bibliographie (Pandoc)
        self.bib_path: Path | None = None
        self.csl_path: Path | None = None
        self.citeproc_enabled: bool = False

        # Widgets
        self.editor = QPlainTextEdit()
        self.preview = PreviewEdit()

        # Typo
        mono = QFont("Consolas")
        mono.setStyleHint(QFont.Monospace)
        mono.setPointSize(11)
        self.editor.setFont(mono)
        self.preview.document().setDefaultFont(QFont("Arial", 11))

        # Split
        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(self.editor)
        splitter.addWidget(self.preview)
        splitter.setSizes([650, 650])
        self.setCentralWidget(splitter)

        # Render suspend pendant édition à droite
        self._suspend_render = False
        self._pending_render = False
        self.preview.on_focus_in = self._preview_focus_in
        self.preview.on_focus_out = self._preview_focus_out

        # Timers : rendu + autosave
        self._render_timer = QTimer(self)
        self._render_timer.setSingleShot(True)
        self._render_timer.setInterval(120)
        self._render_timer.timeout.connect(self._render_preview_now)

        self._autosave_timer = QTimer(self)
        self._autosave_timer.setSingleShot(True)
        self._autosave_timer.setInterval(self.cfg.idle_ms)
        self._autosave_timer.timeout.connect(self._autosave_now)

        self._dirty = False
        self._last_autosaved_hash = None

        self.editor.textChanged.connect(self._on_text_changed)

        # Menus
        self._build_actions()
        self._build_toolbar()
        self.statusBar().showMessage("Prêt")

        # Indicateur Pandoc (widget permanent à droite)
        from PySide6.QtWidgets import QLabel
        self.pandoc_label = QLabel()
        self.statusBar().addPermanentWidget(self.pandoc_label)
        self._update_pandoc_indicator()

        # Contenu initial
        self.editor.setPlainText(
            "# Mini Markdown\n\n"
            "Éditeur à gauche, aperçu à droite.\n\n"
            "- **Gras**, *italique*, `code`\n"
            "- Listes\n"
            "- [lien](https://example.org)\n\n"
            "À droite : tu peux aussi couper/copier/coller (tampon),\n"
            "mais ce que tu y modifies n’est pas réinjecté dans le Markdown.\n"
        )
        self._render_preview_now(force=True)

    def _update_pandoc_indicator(self):
        if self.has_pandoc:
            self.pandoc_label.setText("Pandoc installé")
            self.pandoc_label.setToolTip(self.pandoc_path or "")
        else:
            self.pandoc_label.setText("Pandoc non installé")
            self.pandoc_label.setToolTip("Installe Pandoc ou définis PANDOC_PATH")

    def toggle_citeproc(self, checked: bool):
        self.citeproc_enabled = checked
        if checked and not self.bib_path:
            self.statusBar().showMessage("Citeproc activé, mais aucun .bib sélectionné", 2500)
        else:
            self.statusBar().showMessage(
                "Citations/biblio activées" if checked else "Citations/biblio désactivées",
                1500
            )

    def choose_bib(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Choisir un fichier BibTeX", "", "BibTeX (*.bib);;Tous les fichiers (*)"
        )
        if not path:
            return
        self.bib_path = Path(path)
        self.statusBar().showMessage(f".bib sélectionné : {self.bib_path.name}", 2000)

        # Option pratique : si on choisit un .bib, on active citeproc
        self.citeproc_enabled = True
        if hasattr(self, "act_citeproc"):
            self.act_citeproc.setChecked(True)

    def clear_bib(self):
        self.bib_path = None
        self.statusBar().showMessage(".bib oublié", 1500)

    def choose_csl(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Choisir un fichier CSL (style de citation)", "", "CSL (*.csl);;Tous les fichiers (*)"
        )
        if not path:
            return
        self.csl_path = Path(path)
        self.statusBar().showMessage(f".csl sélectionné : {self.csl_path.name}", 2000)

    def clear_csl(self):
        self.csl_path = None
        self.statusBar().showMessage(".csl oublié", 1500)

    # ---------- UI actions ----------

    def _build_actions(self):
        # Fichier
        act_open = QAction("Ouvrir…", self)
        act_open.setShortcut(QKeySequence.Open)
        act_open.triggered.connect(self.open_file)

        act_save = QAction("Enregistrer", self)
        act_save.setShortcut(QKeySequence.Save)
        act_save.triggered.connect(self.save_file)

        act_save_as = QAction("Enregistrer sous…", self)
        act_save_as.setShortcut(QKeySequence.SaveAs)
        act_save_as.triggered.connect(self.save_file_as)

        act_quit = QAction("Quitter", self)
        act_quit.setShortcut(QKeySequence.Quit)
        act_quit.triggered.connect(self.close)

        # Autosave toggle
        self.act_autosave = QAction("Autosave", self)
        self.act_autosave.setCheckable(True)
        self.act_autosave.setChecked(self.cfg.enabled)
        self.act_autosave.triggered.connect(self.toggle_autosave)

        m_file = self.menuBar().addMenu("Fichier")
        m_file.addAction(act_open)
        m_file.addAction(act_save)
        m_file.addAction(act_save_as)
        m_file.addSeparator()
        m_file.addAction(self.act_autosave)
        m_file.addSeparator()
        m_file.addAction(act_quit)

        # Export
        act_export_html = QAction("Exporter en HTML…", self)
        act_export_html.triggered.connect(self.export_html)

        act_export_pdf = QAction("Exporter en PDF…", self)
        act_export_pdf.triggered.connect(self.export_pdf)

        act_export_docx = QAction("Exporter en DOCX (simple)…", self)
        act_export_docx.triggered.connect(self.export_docx)

        self.act_export_docx_pandoc = QAction("Exporter en DOCX (Pandoc)…", self)
        self.act_export_docx_pandoc.setEnabled(self.has_pandoc)
        self.act_export_docx_pandoc.triggered.connect(self.export_docx_pandoc)

        self.act_export_html_pandoc = QAction("Exporter en HTML (Pandoc)…", self)
        self.act_export_html_pandoc.setEnabled(self.has_pandoc)
        self.act_export_html_pandoc.triggered.connect(self.export_html_pandoc)

        self.act_export_pdf_pandoc = QAction("Exporter en PDF (Pandoc)…", self)
        self.act_export_pdf_pandoc.setEnabled(self.has_pandoc)
        self.act_export_pdf_pandoc.triggered.connect(self.export_pdf_pandoc)

        self.act_export_tex_pandoc = QAction("Exporter en LaTeX (Pandoc)…", self)
        self.act_export_tex_pandoc.setEnabled(self.has_pandoc)
        self.act_export_tex_pandoc.triggered.connect(self.export_tex_pandoc)

        self.act_export_odt_pandoc = QAction("Exporter en ODT (Pandoc)…", self)
        self.act_export_odt_pandoc.setEnabled(self.has_pandoc)
        self.act_export_odt_pandoc.triggered.connect(self.export_odt_pandoc)

        self.act_export_epub_pandoc = QAction("Exporter en EPUB (Pandoc)…", self)
        self.act_export_epub_pandoc.setEnabled(self.has_pandoc)
        self.act_export_epub_pandoc.triggered.connect(self.export_epub_pandoc)


        m_export = self.menuBar().addMenu("Export")
        m_export.addAction(act_export_html)
        m_export.addAction(act_export_pdf)
        m_export.addAction(act_export_docx)
        m_export.addAction(self.act_export_docx_pandoc)
        m_export.addSeparator()
        m_export.addAction(self.act_export_html_pandoc)
        m_export.addAction(self.act_export_pdf_pandoc)
        m_export.addSeparator()
        m_export.addAction(self.act_export_tex_pandoc)
        m_export.addAction(self.act_export_odt_pandoc)
        m_export.addAction(self.act_export_epub_pandoc)

        # Références (Pandoc)
        m_refs = self.menuBar().addMenu("Références")

        self.act_citeproc = QAction("Activer citations + bibliographie (Pandoc)", self)
        self.act_citeproc.setCheckable(True)
        self.act_citeproc.setChecked(self.citeproc_enabled)
        self.act_citeproc.setEnabled(self.has_pandoc)
        self.act_citeproc.triggered.connect(self.toggle_citeproc)

        act_choose_bib = QAction("Choisir un fichier .bib…", self)
        act_choose_bib.setEnabled(self.has_pandoc)
        act_choose_bib.triggered.connect(self.choose_bib)

        act_clear_bib = QAction("Oublier le .bib", self)
        act_clear_bib.setEnabled(self.has_pandoc)
        act_clear_bib.triggered.connect(self.clear_bib)

        act_choose_csl = QAction("Choisir un style .csl… (optionnel)", self)
        act_choose_csl.setEnabled(self.has_pandoc)
        act_choose_csl.triggered.connect(self.choose_csl)

        act_clear_csl = QAction("Oublier le .csl", self)
        act_clear_csl.setEnabled(self.has_pandoc)
        act_clear_csl.triggered.connect(self.clear_csl)

        m_refs.addAction(self.act_citeproc)
        m_refs.addSeparator()
        m_refs.addAction(act_choose_bib)
        m_refs.addAction(act_clear_bib)
        m_refs.addSeparator()
        m_refs.addAction(act_choose_csl)
        m_refs.addAction(act_clear_csl)

        # Édition (agit sur widget focus : gauche ou droite)
        m_edit = self.menuBar().addMenu("Édition")

        act_cut = QAction("Couper", self)
        act_cut.setShortcut(QKeySequence.Cut)
        act_cut.triggered.connect(self._smart_cut)

        act_copy = QAction("Copier", self)
        act_copy.setShortcut(QKeySequence.Copy)
        act_copy.triggered.connect(self._smart_copy)

        act_paste = QAction("Coller", self)
        act_paste.setShortcut(QKeySequence.Paste)
        act_paste.triggered.connect(self._smart_paste)

        m_edit.addAction(act_cut)
        m_edit.addAction(act_copy)
        m_edit.addAction(act_paste)

    def _append_pandoc_citeproc_args(self, cmd: list[str]) -> list[str]:
        """
        Ajoute --citeproc / --bibliography / --csl si activés et valides.
        """
        if self.citeproc_enabled:
            if self.bib_path and self.bib_path.exists():
                cmd.append("--citeproc")
                cmd.append(f"--bibliography={self.bib_path}")
                if self.csl_path and self.csl_path.exists():
                    cmd.append(f"--csl={self.csl_path}")
            else:
                QMessageBox.warning(
                    self, "Bibliographie",
                    "Citeproc est activé, mais aucun fichier .bib valide n’est sélectionné."
                )
        return cmd

    def export_html_pandoc(self):
        if not self.has_pandoc:
            QMessageBox.information(self, "Pandoc", "Pandoc n’est pas disponible.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "Exporter en HTML (Pandoc)", "", "HTML (*.html);;Tous les fichiers (*)"
        )
        if not path:
            return

        out = Path(path)
        if out.suffix.lower() != ".html":
            out = out.with_suffix(".html")

        md = self.editor.toPlainText()

        try:
            import subprocess

            cmd = [
                self.pandoc_path,
                "--from", "markdown",
                "--standalone",
                "--output", str(out),
            ]

            # Optionnel : choisir un moteur PDF (si installé)
            # cmd.append("--pdf-engine=xelatex")  # ou lualatex / pdflatex

            cmd = self._append_pandoc_citeproc_args(cmd)

            subprocess.run(cmd, input=md.encode("utf-8"), check=True)

        except Exception as e:
            QMessageBox.critical(self, "Erreur export HTML (Pandoc)", str(e))
            return

        self.statusBar().showMessage(f"Export HTML (Pandoc) : {out.name}", 1500)

    def export_pdf_pandoc(self):
        if not self.has_pandoc:
            QMessageBox.information(self, "Pandoc", "Pandoc n’est pas disponible.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "Exporter en PDF (Pandoc)", "", "PDF (*.pdf);;Tous les fichiers (*)"
        )
        if not path:
            return

        out = Path(path)
        if out.suffix.lower() != ".pdf":
            out = out.with_suffix(".pdf")

        md = self.editor.toPlainText()

        try:
            import subprocess

            cmd = [
                self.pandoc_path,
                "--from", "markdown",
                "--to", "pdf",
                "--standalone",
                "--output", str(out),
            ]

            # Optionnel : si tu veux forcer un moteur PDF (à installer sur la machine)
            # cmd.append("--pdf-engine=xelatex")

            cmd = self._append_pandoc_citeproc_args(cmd)

            subprocess.run(cmd, input=md.encode("utf-8"), check=True)

        except Exception as e:
            QMessageBox.critical(
                self, "Erreur export PDF (Pandoc)",
                str(e) + "\n\n"
                         "Note : l’export PDF via Pandoc nécessite généralement un moteur LaTeX (TeX Live / MiKTeX)."
            )
            return

        self.statusBar().showMessage(f"Export PDF (Pandoc) : {out.name}", 1500)

    def export_with_pandoc(self, to_format: str, dialog_title: str, filter_str: str, default_suffix: str):
        if not self.has_pandoc:
            QMessageBox.information(self, "Pandoc", "Pandoc n’est pas disponible.")
            return

        path, _ = QFileDialog.getSaveFileName(self, dialog_title, "", filter_str)
        if not path:
            return

        out = Path(path)
        if out.suffix.lower() != default_suffix:
            out = out.with_suffix(default_suffix)

        md = self.editor.toPlainText()

        try:
            import subprocess

            cmd = [
                self.pandoc_path,
                "--from", "markdown",
                "--to", to_format,
                "--standalone",
                "--output", str(out),
            ]

            # Biblio/citations si activé
            cmd = self._append_pandoc_citeproc_args(cmd)

            # Styles Word optionnels pour DOCX seulement
            if to_format == "docx" and self.current_path:
                ref = self.current_path.parent / "reference.docx"
                if ref.exists():
                    cmd.append(f"--reference-doc={ref}")

            subprocess.run(cmd, input=md.encode("utf-8"), check=True)

        except Exception as e:
            QMessageBox.critical(self, f"Erreur export ({to_format})", str(e))
            return

        self.statusBar().showMessage(f"Export Pandoc ({to_format}) : {out.name}", 1500)

    def export_tex_pandoc(self):
        self.export_with_pandoc(
            to_format="latex",
            dialog_title="Exporter en LaTeX (Pandoc)",
            filter_str="LaTeX (*.tex);;Tous les fichiers (*)",
            default_suffix=".tex",
        )

    def export_odt_pandoc(self):
        self.export_with_pandoc(
            to_format="odt",
            dialog_title="Exporter en ODT (Pandoc)",
            filter_str="ODT (*.odt);;Tous les fichiers (*)",
            default_suffix=".odt",
        )

    def export_epub_pandoc(self):
        self.export_with_pandoc(
            to_format="epub",
            dialog_title="Exporter en EPUB (Pandoc)",
            filter_str="EPUB (*.epub);;Tous les fichiers (*)",
            default_suffix=".epub",
        )

    def export_docx_pandoc(self):
        if not self.has_pandoc:
            QMessageBox.information(
                self, "Pandoc",
                "Pandoc n’est pas disponible.\n\n"
                "Installe Pandoc (ou ajoute-le au PATH), ou définis PANDOC_PATH."
            )
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "Exporter en DOCX (Pandoc)", "", "Word (*.docx);;Tous les fichiers (*)"
        )
        if not path:
            return

        out = Path(path)
        if out.suffix.lower() != ".docx":
            out = out.with_suffix(".docx")

        md = self.editor.toPlainText()

        try:
            import subprocess

            cmd = [
                self.pandoc_path,
                "--from", "markdown",
                "--to", "docx",
                "--output", str(out),
                "--standalone",
            ]

            # Option : styles Word si reference.docx à côté du .md
            if self.current_path:
                ref = self.current_path.parent / "reference.docx"
                if ref.exists():
                    cmd.append(f"--reference-doc={ref}")

            # Bibliographie / citations (Pandoc)
            cmd = self._append_pandoc_citeproc_args(cmd)

            subprocess.run(cmd, input=md.encode("utf-8"), check=True)

        except Exception as e:
            QMessageBox.critical(self, "Erreur export DOCX (Pandoc)", str(e))
            return

        self.statusBar().showMessage(f"Export DOCX (Pandoc) : {out.name}", 1500)


    def _find_pandoc(self) -> str | None:
        """
        Détecte pandoc :
        - variable d’environnement PANDOC_PATH
        - PATH
        """
        env = os.environ.get("PANDOC_PATH")
        if env and Path(env).exists():
            return env

        which = shutil.which("pandoc")
        if which:
            return which

        return None

    def _build_toolbar(self):
        tb = self.addToolBar("Mise en forme")
        tb.setMovable(False)

        # Boutons principaux
        act_bold = QAction("B", self)
        act_bold.setToolTip("Gras (**...**)  | Ctrl+B")
        act_bold.setShortcut(QKeySequence.Bold)  # Ctrl+B
        act_bold.triggered.connect(lambda: self._wrap_selection("**", "**"))
        self.addAction(act_bold)
        tb.addAction(act_bold)

        act_italic = QAction("I", self)
        act_italic.setToolTip("Italique (*...*)  | Ctrl+I")
        act_italic.setShortcut(QKeySequence.Italic)  # Ctrl+I
        act_italic.triggered.connect(lambda: self._wrap_selection("*", "*"))
        self.addAction(act_italic)
        tb.addAction(act_italic)

        act_underline = QAction("U", self)
        act_underline.setToolTip("Souligné (<u>...</u>)  | Ctrl+U")
        act_underline.setShortcut(QKeySequence.Underline)  # Ctrl+U
        act_underline.triggered.connect(lambda: self._wrap_selection("<u>", "</u>"))
        self.addAction(act_underline)
        tb.addAction(act_underline)

        tb.addSeparator()

        act_code = QAction("`code`", self)
        act_code.setToolTip("Code (`...`) ou bloc ```...```")
        act_code.triggered.connect(self._toggle_code)
        tb.addAction(act_code)

        act_h1 = QAction("H1", self)
        act_h1.setToolTip("Titre (préfixe #)")
        act_h1.triggered.connect(lambda: self._prefix_lines("# "))
        tb.addAction(act_h1)

        act_h2 = QAction("H2", self)
        act_h2.setToolTip("Sous-titre (préfixe ##)")
        act_h2.triggered.connect(lambda: self._prefix_lines("## "))
        tb.addAction(act_h2)

        tb.addSeparator()

        act_ul = QAction("•", self)
        act_ul.setToolTip("Liste à puces (préfixe - )")
        act_ul.triggered.connect(lambda: self._prefix_lines("- "))
        tb.addAction(act_ul)

        act_ol = QAction("1.", self)
        act_ol.setToolTip("Liste numérotée (préfixe 1. )")
        act_ol.triggered.connect(lambda: self._prefix_lines("1. "))
        tb.addAction(act_ol)

        tb.addSeparator()

        act_link = QAction("Lien", self)
        act_link.setToolTip("Lien [texte](url)")
        act_link.triggered.connect(self._insert_link)
        tb.addAction(act_link)

        tb.addSeparator()

        act_quote = QAction("❝", self)
        act_quote.setToolTip("Blockquote (préfixe > )")
        act_quote.triggered.connect(lambda: self._prefix_lines("> "))
        tb.addAction(act_quote)

        act_hr = QAction("—", self)
        act_hr.setToolTip("Ligne horizontale (---)")
        act_hr.triggered.connect(self._insert_hr)
        tb.addAction(act_hr)

        act_table = QAction("Table", self)
        act_table.setToolTip("Insérer un gabarit de table Markdown")
        act_table.triggered.connect(self._insert_table)
        tb.addAction(act_table)

        act_img = QAction("Image", self)
        act_img.setToolTip("Image ![alt](url)")
        act_img.triggered.connect(self._insert_image)
        tb.addAction(act_img)


    def _active_editor(self) -> QPlainTextEdit:
        """
        On applique la mise en forme au Markdown (gauche), jamais au tampon (droite).
        """
        return self.editor

    def _wrap_selection(self, left: str, right: str):
        ed = self._active_editor()
        cursor = ed.textCursor()

        if cursor.hasSelection():
            selected = cursor.selectedText()
            # Qt met parfois des séparateurs Unicode pour les retours ligne en sélection
            selected = selected.replace("\u2029", "\n")
            cursor.insertText(f"{left}{selected}{right}")
        else:
            # Insère les marqueurs et place le curseur au milieu
            cursor.insertText(f"{left}{right}")
            cursor.movePosition(QTextCursor.Left, QTextCursor.MoveAnchor, len(right))
            ed.setTextCursor(cursor)

        ed.setFocus()

    def _prefix_lines(self, prefix: str):
        """
        Préfixe la/les ligne(s) sélectionnée(s) par `prefix`.
        Si pas de sélection : agit sur la ligne courante.
        """
        ed = self._active_editor()
        cursor = ed.textCursor()

        # Étendre à des lignes entières
        start = cursor.selectionStart()
        end = cursor.selectionEnd()
        cursor.setPosition(start)
        cursor.movePosition(QTextCursor.StartOfLine)
        start_line = cursor.position()

        cursor.setPosition(end)
        cursor.movePosition(QTextCursor.EndOfLine)
        end_line = cursor.position()

        cursor.setPosition(start_line)
        cursor.setPosition(end_line, QTextCursor.KeepAnchor)

        block = cursor.selectedText().replace("\u2029", "\n")
        lines = block.split("\n")
        lines = [prefix + ln if ln.strip() else ln for ln in lines]
        cursor.insertText("\n".join(lines))
        ed.setFocus()

    def _toggle_code(self):
        """
        Si sélection sur une seule ligne : `inline code`
        Si sélection multi-lignes : bloc ``` ```
        Si pas de sélection : insère `` et place le curseur au milieu.
        """
        ed = self._active_editor()
        cursor = ed.textCursor()

        if not cursor.hasSelection():
            self._wrap_selection("`", "`")
            return

        selected = cursor.selectedText().replace("\u2029", "\n")
        if "\n" in selected:
            cursor.insertText(f"```\n{selected}\n```")
        else:
            cursor.insertText(f"`{selected}`")
        ed.setFocus()

    def _insert_hr(self):
        """
        Insère une ligne horizontale Markdown, sur une ligne isolée.
        """
        ed = self._active_editor()
        cursor = ed.textCursor()
        cursor.beginEditBlock()

        cursor.movePosition(QTextCursor.EndOfLine)
        cursor.insertText("\n\n---\n\n")

        cursor.endEditBlock()
        ed.setTextCursor(cursor)
        ed.setFocus()

    def _insert_table(self):
        """
        Insère un gabarit simple de table Markdown (2 colonnes).
        Si une sélection existe et contient des lignes -> tente de convertir en table 2 colonnes via séparation par tab.
        Sinon -> insère un modèle.
        """
        ed = self._active_editor()
        cursor = ed.textCursor()

        if cursor.hasSelection():
            selected = cursor.selectedText().replace("\u2029", "\n").strip("\n")
            lines = [ln for ln in selected.split("\n") if ln.strip()]
            # Tentative: chaque ligne "col1\tcol2"
            rows = []
            ok = True
            for ln in lines:
                if "\t" not in ln:
                    ok = False
                    break
                a, b = ln.split("\t", 1)
                rows.append((a.strip(), b.strip()))

            if ok and rows:
                header = "| Colonne 1 | Colonne 2 |\n|---|---|\n"
                body = "\n".join([f"| {a} | {b} |" for a, b in rows]) + "\n"
                cursor.insertText(header + body)
                ed.setFocus()
                return

        # Modèle par défaut
        template = (
            "\n\n"
            "| Colonne 1 | Colonne 2 |\n"
            "|---|---|\n"
            "| Valeur 1 | Valeur 2 |\n"
            "| Valeur 3 | Valeur 4 |\n"
            "\n"
        )
        cursor.insertText(template)
        ed.setTextCursor(cursor)
        ed.setFocus()

    def _insert_image(self):
        """
        Insère une image Markdown.
        - Si sélection : ![selection](url)
        - Sinon : ![alt](url) avec curseur sur alt.
        """
        ed = self._active_editor()
        cursor = ed.textCursor()

        if cursor.hasSelection():
            alt = cursor.selectedText().replace("\u2029", "\n")
            cursor.insertText(f"![{alt}](https://)")
            cursor.movePosition(QTextCursor.Left, QTextCursor.MoveAnchor, 1)
            ed.setTextCursor(cursor)
        else:
            cursor.insertText("![alt](https://)")
            # placer le curseur sur "alt"
            cursor.movePosition(QTextCursor.Left, QTextCursor.MoveAnchor, len("](https://)"))
            cursor.movePosition(QTextCursor.Left, QTextCursor.MoveAnchor, len("alt"))
            ed.setTextCursor(cursor)

        ed.setFocus()

    def _insert_link(self):
        """
        Insère un lien Markdown.
        - Si sélection : [selection](url)
        - Sinon : [texte](url) avec curseur sur 'texte'
        """
        ed = self._active_editor()
        cursor = ed.textCursor()

        if cursor.hasSelection():
            text = cursor.selectedText().replace("\u2029", "\n")
            cursor.insertText(f"[{text}](https://)")
            # placer le curseur après https:// pour saisir l'URL
            cursor.movePosition(QTextCursor.Left, QTextCursor.MoveAnchor, 1)
            ed.setTextCursor(cursor)
        else:
            cursor.insertText("[texte](https://)")
            # placer le curseur sur "texte"
            cursor.movePosition(QTextCursor.Left, QTextCursor.MoveAnchor, len("](https://)"))
            cursor.movePosition(QTextCursor.Left, QTextCursor.MoveAnchor, len("texte"))
            ed.setTextCursor(cursor)

        ed.setFocus()

    def toggle_autosave(self, checked: bool):
        self.cfg.enabled = checked
        if not checked:
            self._autosave_timer.stop()
            self.statusBar().showMessage("Autosave désactivé", 1200)
        else:
            self.statusBar().showMessage("Autosave activé", 1200)
            if self._dirty:
                self._autosave_timer.start(self.cfg.idle_ms)

    def _smart_cut(self):
        w = QApplication.focusWidget()
        if hasattr(w, "cut"):
            w.cut()

    def _smart_copy(self):
        w = QApplication.focusWidget()
        if hasattr(w, "copy"):
            w.copy()

    def _smart_paste(self):
        w = QApplication.focusWidget()
        if hasattr(w, "paste"):
            w.paste()

    # ---------- Rendering + focus ----------

    def _preview_focus_in(self):
        self._suspend_render = True
        self.statusBar().showMessage("Aperçu : édition active (rendu suspendu)", 1200)

    def _preview_focus_out(self):
        self._suspend_render = False
        if self._pending_render:
            self._pending_render = False
            self._render_preview_now(force=True)

    def _on_text_changed(self):
        self._dirty = True

        # rendu
        if self._suspend_render:
            self._pending_render = True
        else:
            self._render_timer.start()

        # autosave
        if self.cfg.enabled:
            self._autosave_timer.start(self.cfg.idle_ms)

    def _render_preview_now(self, force: bool = False):
        if self._suspend_render and not force:
            self._pending_render = True
            return
        md = self.editor.toPlainText()
        # On met à jour la preview (écrase le “tampon” si focus à gauche)
        self.preview.blockSignals(True)
        self.preview.setMarkdown(md)
        self.preview.blockSignals(False)

    # ---------- File ops ----------

    def open_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Ouvrir un fichier Markdown", "", "Markdown (*.md *.markdown);;Tous les fichiers (*)"
        )
        if not path:
            return
        p = Path(path)
        try:
            text = p.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = p.read_text(encoding="utf-8", errors="replace")

        self.editor.setPlainText(text)
        self.current_path = p
        self._dirty = False
        self._last_autosaved_hash = None

        self.statusBar().showMessage(f"Ouvert : {p.name}", 1500)
        self._render_preview_now(force=True)

    def save_file(self):
        if self.current_path is None:
            self.save_file_as()
            return
        try:
            self.current_path.write_text(self.editor.toPlainText(), encoding="utf-8")
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Impossible d’enregistrer :\n{e}")
            return
        self._dirty = False
        self.statusBar().showMessage(f"Enregistré : {self.current_path.name}", 1500)

    def save_file_as(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "Enregistrer sous", "", "Markdown (*.md);;Tous les fichiers (*)"
        )
        if not path:
            return
        p = Path(path)
        if p.suffix.lower() not in (".md", ".markdown"):
            p = p.with_suffix(".md")
        self.current_path = p
        self.save_file()

    # ---------- Autosave ----------

    def _autosave_target(self) -> Path:
        """
        Si on a un fichier courant : autosave dessus (comme demandé “en temps réel”).
        Sinon : fallback dans le HOME.
        """
        if self.current_path and self.cfg.use_main_file_if_possible:
            return self.current_path
        return Path.home() / self.cfg.fallback_filename

    def _autosave_now(self):
        if not self.cfg.enabled:
            return

        md = self.editor.toPlainText()
        h = hash(md)
        if h == self._last_autosaved_hash:
            return

        target = self._autosave_target()
        try:
            target.write_text(md, encoding="utf-8")
        except Exception as e:
            # On ne spam pas de popups : juste un message bref
            self.statusBar().showMessage(f"Autosave échoué : {e}", 2500)
            return

        self._last_autosaved_hash = h
        self.statusBar().showMessage(f"Autosave : {target.name}", 900)

    # ---------- Exports ----------

    def export_html(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "Exporter en HTML", "", "HTML (*.html);;Tous les fichiers (*)"
        )
        if not path:
            return
        out = Path(path)
        if out.suffix.lower() != ".html":
            out = out.with_suffix(".html")

        md = self.editor.toPlainText()
        doc = QTextDocument()
        doc.setMarkdown(md)
        html = doc.toHtml()

        try:
            out.write_text(html, encoding="utf-8")
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Impossible d’exporter HTML :\n{e}")
            return
        self.statusBar().showMessage(f"Export HTML : {out.name}", 1500)

    def export_pdf(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "Exporter en PDF", "", "PDF (*.pdf);;Tous les fichiers (*)"
        )
        if not path:
            return
        out = Path(path)
        if out.suffix.lower() != ".pdf":
            out = out.with_suffix(".pdf")

        md = self.editor.toPlainText()
        doc = QTextDocument()
        doc.setMarkdown(md)

        printer = QPrinter(QPrinter.HighResolution)
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(str(out))

        try:
            doc.print_(printer)
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Impossible d’exporter PDF :\n{e}")
            return

        self.statusBar().showMessage(f"Export PDF : {out.name}", 1500)

    def export_docx(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "Exporter en DOCX", "", "Word (*.docx);;Tous les fichiers (*)"
        )
        if not path:
            return
        out = Path(path)
        if out.suffix.lower() != ".docx":
            out = out.with_suffix(".docx")

        md = self.editor.toPlainText()
        try:
            export_docx_from_markdown(md, out)
        except ModuleNotFoundError:
            QMessageBox.critical(
                self, "DOCX",
                "Il manque la dépendance python-docx.\n\nInstalle : pip install python-docx"
            )
            return
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Impossible d’exporter DOCX :\n{e}")
            return

        self.statusBar().showMessage(f"Export DOCX : {out.name}", 1500)


def main():
    app = QApplication(sys.argv)
    win = MainWindow()
    win.resize(1300, 780)
    win.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
